from docx import Document
import os

def combine_txt_to_docx(txt_folder='output', output_filename='final_translation.docx'):
    doc = Document()
    part_files = sorted(
        [f for f in os.listdir(txt_folder) if f.endswith('.txt')],
        key=lambda x: int(''.join(filter(str.isdigit, x)))  # Sort by Part number
    )

    for file in part_files:
        filepath = os.path.join(txt_folder, file)
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            doc.add_heading(file.replace('.txt', ''), level=2)
            doc.add_paragraph(content)
            doc.add_page_break()

    final_path = os.path.join(txt_folder, output_filename)
    doc.save(final_path)
    return final_path
